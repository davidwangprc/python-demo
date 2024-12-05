import pandas as pd
import json
from docx import Document
import os

def excel_to_json(excel_path, json_path):
    try:
        # 读取Excel文件
        df = pd.read_excel(excel_path)
        
        # 将DataFrame转换为字典列表
        result = []
        for _, row in df.iterrows():
            item = {
                "college": row["学院"],
                "major": row["专业"],
                "graduation_requirement": row["毕业要求-1"]
            }
            result.append(item)
            
        # 写入JSON文件
        with open(json_path, 'w', encoding='utf-8') as f:
            json.dump(result, f, ensure_ascii=False, indent=2)
            
        print(f"转换成功！JSON文件已保存至: {json_path}")
        
    except Exception as e:
        print(f"转换过程中出现错误: {str(e)}")

def create_word_from_template(data, template_path, output_dir='output'):
    # 确保主输出目录存在
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)
        
    # 为每个专业创建文档
    for item in data:
        try:
            # 读取模板
            doc = Document(template_path)
            
            # 替换所有段落中的占位符
            for paragraph in doc.paragraphs:
                if '{目标}' in paragraph.text:
                    paragraph.text = paragraph.text.replace('{目标}', item['graduation_requirement'])
            
            # 替换表格中的占位符
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if '{目标}' in cell.text:
                            cell.text = cell.text.replace('{目标}', item['graduation_requirement'])
            
            # 创建学院子文件夹
            college_dir = os.path.join(output_dir, item['college'])
            if not os.path.exists(college_dir):
                os.makedirs(college_dir)
            
            # 设置文件名并保存
            prefix = '2024-2025-1'
            suffix = '教学大纲-中国近现代史纲要'
            filename = f"{prefix}-{item['major']}-{suffix}.docx"
            # 替换可能导致问题的字符
            filename = filename.replace(':', '_').replace('/', '_').replace('\\', '_')
            
            # 使用学院子文件夹路径
            filepath = os.path.join(college_dir, filename)
            doc.save(filepath)
            print(f"已生成文档: {filepath}")
            
        except Exception as e:
            print(f"处理 {item['college']}-{item['major']} 时出错: {str(e)}")

def main():
    # 文件路径
    input_excel = 'data/专业对应目标提取表.xlsx'
    temp_json = 'data/专业对应目标提取表.json'
    template_path = 'data/template_demo.docx'
    
    # 首先转换Excel到JSON
    excel_to_json(input_excel, temp_json)
    
    # 读取JSON数据
    with open(temp_json, 'r', encoding='utf-8') as f:
        data = json.load(f)
    
    # 根据模板创建Word文档
    create_word_from_template(data, template_path)

if __name__ == "__main__":
    main()